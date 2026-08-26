from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "model_dataset.csv"
WORK_DIR = ROOT / ".docx_work"
HEATMAP_PATH = WORK_DIR / "target_correlation_heatmap.png"
OUTPUT_PATH = ROOT / "Project_Phase2_Answers.docx"

# Resolved design preset: standard_business_brief.
# Named overrides:
# - academic_title: 22 pt dark navy, used once in the opening title block.
# - dense_explanation_table: 9.25 pt body / 9.5 pt header so ten detailed rows
#   remain readable on one page without fixed row heights.
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "222222"
MUTED = "5A6573"
LIGHT_FILL = "F2F4F7"
BORDER = "B8C2CC"
WHITE = "FFFFFF"

FEATURES = [
    (
        "elo_diff",
        "Elo rating difference",
        "A higher pre-game home-team Elo relative to the away team is associated with a greater probability of a home win.",
    ),
    (
        "decay_goal_diff_diff",
        "Decayed goal-difference gap",
        "Better recent scoring margin for the home team, with older games given less weight, is positively related to a home win.",
    ),
    (
        "last_10_es_shot_attempt_share_diff",
        "Even-strength shot-attempt share gap",
        "A home advantage in recent even-strength possession is positively related to winning.",
    ),
    (
        "decay_win_rate_diff",
        "Decayed win-rate gap",
        "A stronger recency-weighted win rate for the home team is positively related to the target.",
    ),
    (
        "last_10_es_shot_attempts_against_avg_diff",
        "Even-strength attempts-against gap",
        "The negative sign means the home team is less likely to win when it has recently allowed more even-strength attempts than the away team.",
    ),
    (
        "last_10_shots_for_avg_diff",
        "Shots-for gap (last 10)",
        "Generating more recent shots than the opponent is positively related to a home win.",
    ),
    (
        "season_points_pct_diff",
        "Season points-percentage gap",
        "A higher pre-game season points percentage for the home team is positively related to winning.",
    ),
    (
        "last_10_shots_against_avg_diff",
        "Shots-against gap (last 10)",
        "The negative relationship indicates that allowing more shots than the away team reduces the home team's win tendency.",
    ),
    (
        "last_10_es_shot_attempts_for_avg_diff",
        "Even-strength attempts-for gap",
        "A greater recent volume of even-strength attempts by the home team is positively related to the target.",
    ),
    (
        "goal_diff_last_10_diff",
        "Goal-difference gap (last 10)",
        "A better recent goal difference for the home team is positively related to a home win.",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def apply_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_body_paragraph(doc, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def set_image_alt_text(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def build_heatmap(df: pd.DataFrame) -> pd.Series:
    correlations = df[[key for key, _, _ in FEATURES]].corrwith(df["target_home_win"])
    values = np.array([correlations[key] for key, _, _ in FEATURES], dtype=float).reshape(-1, 1)
    labels = [label for _, label, _ in FEATURES]

    sns.set_theme(style="white", font_scale=0.9)
    fig, ax = plt.subplots(figsize=(7.0, 5.7), dpi=220)
    sns.heatmap(
        values,
        ax=ax,
        cmap="RdBu",
        vmin=-0.20,
        vmax=0.20,
        center=0,
        annot=np.array([[f"{value:+.3f}"] for value in values[:, 0]]),
        fmt="",
        linewidths=0.8,
        linecolor="white",
        yticklabels=labels,
        xticklabels=["target_home_win"],
        cbar_kws={"label": "Pearson correlation (r)", "shrink": 0.86},
    )
    ax.set_title("Selected predictor correlations with the home-win target", fontsize=13, weight="bold", pad=12)
    ax.set_xlabel("")
    ax.set_ylabel("")
    ax.tick_params(axis="y", labelsize=9.2, rotation=0)
    ax.tick_params(axis="x", labelsize=9.5, rotation=0)
    fig.tight_layout()
    fig.savefig(HEATMAP_PATH, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return correlations


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.0

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("PROJECT PHASE 2  •  DATA PREPARATION")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("Project Phase 2 Answers  |  Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(fp)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Project Phase 2 — Question Answers")
    set_run_font(run, size=22, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True
    run = subtitle.add_run("Data Wrangling, Feature Selection and Train–Test Split | NHL winner-prediction dataset")
    set_run_font(run, size=11, color=MUTED, italic=True)


def add_question_1(doc: Document) -> None:
    doc.add_heading("1. What is the Pearson correlation coefficient?", level=1)
    add_body_paragraph(
        doc,
        "The Pearson correlation coefficient, written as r, measures the direction and strength of the linear relationship between two numerical variables. It compares how the variables vary together with how much they vary separately.",
    )

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(8)
    formula.paragraph_format.space_after = Pt(10)
    formula.paragraph_format.keep_together = True
    run = formula.add_run("r = covariance(X, Y) / (sₓ × sᵧ)")
    set_run_font(run, size=12, color=NAVY, italic=True)

    add_body_paragraph(
        doc,
        "Its value ranges from −1 to +1. A value near +1 indicates a strong positive linear relationship, a value near −1 indicates a strong negative linear relationship, and a value near 0 indicates little or no linear relationship. The magnitude shows the strength and the sign shows the direction.",
    )
    add_body_paragraph(
        doc,
        "For this project, the target is binary: target_home_win equals 1 when the home team wins and 0 otherwise. Pearson correlation with a binary variable is equivalent to the point-biserial correlation, so a positive r means higher feature values are associated with home wins, while a negative r means higher values are associated with away wins. Correlation does not prove causation, and it may miss nonlinear relationships or interactions between predictors.",
    )


def add_question_2(doc: Document, df: pd.DataFrame, correlations: pd.Series) -> None:
    doc.add_page_break()
    doc.add_heading("2. How is each attribute important for predicting the target (heat map)?", level=1)
    add_body_paragraph(
        doc,
        "The heat map uses 20,591 completed NHL games and target_home_win (1 = home win, 0 = away win). It shows ten nonredundant, comparison-style predictors with the largest absolute Pearson relationships to the target. Raw home/away components and deterministic transformations were omitted where they repeated the same signal. Correlations were calculated with the available nonmissing pairs for each feature.",
    )

    picture_paragraph = doc.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_before = Pt(4)
    picture_paragraph.paragraph_format.space_after = Pt(0)
    run = picture_paragraph.add_run()
    shape = run.add_picture(str(HEATMAP_PATH), width=Inches(5.85))
    set_image_alt_text(
        shape,
        "Target correlation heat map",
        "A one-column heat map of ten NHL predictors and their Pearson correlations with target_home_win. Positive values are blue and negative values are red.",
    )

    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run("Figure 1. Pearson correlations with target_home_win. Blue is positive; red is negative.")

    add_body_paragraph(
        doc,
        "The strongest displayed relationship is Elo rating difference (r = +0.179). All coefficients are modest, which means no single predictor is sufficient by itself; the prediction model should combine several features and can benefit from algorithms that capture nonlinear effects and interactions.",
    )

    doc.add_page_break()
    heading = doc.add_heading("Attribute-by-attribute interpretation", level=2)
    keep_with_next(heading)
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.0
    run = source.add_run("The sign is interpreted from the home team’s perspective; larger absolute values indicate stronger linear screening value.")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    keep_with_next(source)

    table = doc.add_table(rows=1, cols=3)
    headers = ["Attribute", "r", "Relationship to target_home_win"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, color=NAVY, bold=True)

    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

    for key, label, explanation in FEATURES:
        cells = table.add_row().cells
        values = [label, f"{correlations[key]:+.3f}", explanation]
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.02
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            set_run_font(run, size=9.25, color=INK, bold=(index == 0))

    apply_table_geometry(table, [2550, 720, 6090])
    set_table_borders(table)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    run = note.add_run(
        "These correlations are useful for feature screening, but they are not complete model-based importance scores. Final importance should be confirmed inside a leakage-safe validation process."
    )
    set_run_font(run, size=9.5, color=MUTED, italic=True)


def add_question_3(doc: Document) -> None:
    doc.add_heading("3. What is k-fold cross-validation?", level=1)
    add_body_paragraph(
        doc,
        "K-fold cross-validation is a resampling method used to estimate how well a model will perform on unseen data. The available training data is divided into k approximately equal groups called folds. The model is trained k times. On each run, one fold is used for validation and the other k − 1 folds are used for training. Every observation is therefore used for validation once and for training k − 1 times.",
    )
    add_body_paragraph(
        doc,
        "The k validation scores are averaged to produce the cross-validation score, and their variation shows how stable the model is across different samples. For example, in 5-fold cross-validation, the model trains on about 80% of the development data and validates on about 20% in each of five runs. Common choices are k = 5 or k = 10 because they provide a useful balance between reliability and computation time.",
    )
    add_body_paragraph(
        doc,
        "For classification, stratified k-fold is often used so each fold has a similar proportion of home wins and away wins. All learned preprocessing—including imputation, scaling, feature selection and tuning—must be fitted only on the training portion of each fold; otherwise, information from the validation fold leaks into the model and makes the result too optimistic.",
    )
    add_body_paragraph(
        doc,
        "Because NHL games are chronological, randomly shuffled k-fold validation is not appropriate for the final model evaluation. A time-series or rolling-origin version should train on earlier games and validate on later games. This preserves the real prediction situation and prevents future results from influencing past predictions.",
    )


def add_question_4(doc: Document) -> None:
    doc.add_heading("4. Why is the training dataset 70–80%? Why is the test dataset 20–30%?", level=1)
    add_body_paragraph(
        doc,
        "A 70–80% training share gives the model most of the available observations so it can learn stable patterns, estimate parameters and represent less common situations. If the training set is too small, the fitted model can have high variance and may not learn the relationship between the predictors and the target reliably.",
    )
    add_body_paragraph(
        doc,
        "A 20–30% test share reserves enough completely unseen observations to measure generalization with useful statistical precision. If the test set is too small, performance estimates can change greatly because of a few games. If it is too large, model quality may decrease because too much information was withheld from training.",
    )
    add_body_paragraph(
        doc,
        "Therefore, 70/30 and 80/20 are practical compromises rather than strict rules. An 80/20 split is common when the dataset is large and the model benefits from additional training data; 70/30 provides a larger independent test sample when a more precise evaluation is needed. With 20,591 games, these choices correspond to approximately 16,473 training and 4,118 test games for 80/20, or 14,414 training and 6,177 test games for 70/30.",
    )
    add_body_paragraph(
        doc,
        "The test set must remain untouched until the final evaluation. Cross-validation is performed only within the training portion for model selection. For this NHL project, the split should also respect time: earlier seasons and games belong in training, while the most recent period belongs in testing. This gives a more realistic estimate of performance on future games than a random split.",
    )


def main() -> None:
    WORK_DIR.mkdir(exist_ok=True)
    df = pd.read_csv(DATA_PATH, low_memory=False)
    correlations = build_heatmap(df)

    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_question_1(doc)
    add_question_2(doc, df, correlations)
    add_question_3(doc)
    add_question_4(doc)

    core = doc.core_properties
    core.title = "Project Phase 2 — Question Answers"
    core.subject = "Pearson correlation, attribute relationships, k-fold cross-validation and train-test split"
    core.author = ""
    core.keywords = "data preparation, Pearson correlation, heat map, cross-validation, train-test split, NHL"
    core.comments = "Prepared as answers to the four questions in Project Phase 2."

    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")
    print(correlations.to_string())


if __name__ == "__main__":
    main()
